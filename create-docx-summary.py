from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
import io
import os
import re

print("Generating report...")

# ---------- helpers ----------
def resize_image_to_limit(filepath, max_bytes=512000, min_quality=10):
    with Image.open(filepath) as img:
        quality = 95
        while quality >= min_quality:
            buffer = io.BytesIO()
            img.save(buffer, format='JPEG', quality=quality)
            if buffer.tell() <= max_bytes:
                with open(filepath, 'wb') as f:
                    f.write(buffer.getvalue())
                return True
            quality -= 5
    return False

def extract_slice_number_from_filename(filename):
    # grabs the first integer in the filename
    m = re.search(r"(\d+)", filename)
    return int(m.group(1)) if m else None

def collect_slice_numbers(folder, pattern_prefix=None):
    """Return sorted set of slice numbers for files in 'folder' that end with .bmp.
       Optionally ensure filenames start with pattern_prefix."""
    nums = set()
    if not os.path.isdir(folder):
        return nums
    for f in os.listdir(folder):
        if not f.lower().endswith(".bmp"):
            continue
        if pattern_prefix and not f.startswith(pattern_prefix):
            continue
        n = extract_slice_number_from_filename(f)
        if n is not None:
            nums.add(n)
    return nums

# ---------- doc builder ----------
def generate_report(axis="x"):
    axis = axis.lower()
    if axis not in ("x", "y", "z"):
        print("Invalid axis. Use 'x', 'y', or 'z'.")
        return

    script_dir = os.path.dirname(os.path.abspath(__file__))

    folders = {
        "disp":  os.path.join(script_dir, "exports", "displacements", f"{axis}slice"),
        "max":   os.path.join(script_dir, "exports", "max_principal", f"{axis}slice"),
        "min":   os.path.join(script_dir, "exports", "min_principal", f"{axis}slice"),
        "state": os.path.join(script_dir, "exports", "zone_state",    f"{axis}slice"),
        "zz":    os.path.join(script_dir, "exports", "zz_stress",     f"{axis}slice"),  # NEW
    }

    filename_templates = {
        "disp":  f"{axis}_slice_disp_{{}}.bmp",
        "max":   f"{axis}_slice_max_principal_{{}}.bmp",
        "min":   f"{axis}_slice_min_principal_{{}}.bmp",
        "state": f"{axis}_slice_state_{{}}.bmp",
        "zz":    f"{axis}_slice_zz_{{}}.bmp",  # NEW
    }

    labels = {
        "disp":  "Displacement Magnitude",
        "max":   "Max Principal Effective Stress",
        "min":   "Min Principal Effective Stress",
        "state": "Zone State",
        "zz":    "σzz Effective Stress",
    }

    # Page layout: two per page, then zz on its own page
    pages = [
        ("disp", "max"),
        ("min",  "state"),
        ("zz",),
    ]

    # Collect slice numbers PER KEY, then take union so we don’t assume all keys share the same set
    # Optionally guard with a prefix to reduce false matches
    per_key_numbers = {
        k: collect_slice_numbers(folders[k])
        for k in folders
    }
    all_slice_numbers = sorted(set().union(*per_key_numbers.values()))
    if not all_slice_numbers:
        print("No slice images found for any key; nothing to write.")
        return

    # Build the doc
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.7)

    # Slim captions
    if "Caption" in document.styles:
        cap = document.styles["Caption"]
        cap.font.size = Pt(9)
        cap.paragraph_format.space_before = Pt(0)
        cap.paragraph_format.space_after = Pt(6)

    if "Normal" in document.styles:
        normal = document.styles["Normal"]
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)

    IMG_W = Inches(5.5)
    IMG_H = Inches(4.0)

    total_pages = len(pages)

    for slice_val in all_slice_numbers:
        slice_str = str(slice_val)

        for page_idx, keys in enumerate(pages, start=1):
            # Check if at least one image exists for this page/slice; skip page if not
            existing = []
            for key in keys:
                bmp_name = filename_templates[key].format(slice_str)
                bmp_path = os.path.join(folders[key], bmp_name)
                if os.path.exists(bmp_path):
                    existing.append((key, bmp_path))

            if not existing:
                # nothing to show for this page for this slice; skip
                continue

            document.add_heading(f"{axis.upper()} Slice @ {slice_str}  ({page_idx}/{total_pages})", level=1)

            table = document.add_table(rows=0, cols=1)
            table.autofit = False

            for key, bmp_path in existing:
                # Convert BMP -> JPG (to shrink doc) then insert
                jpg_path = bmp_path.replace(".bmp", ".jpg")
                with Image.open(bmp_path) as im:
                    im.convert("RGB").save(jpg_path, "JPEG", quality=95)
                resize_image_to_limit(jpg_path)

                row_img = table.add_row().cells[0]
                run = row_img.paragraphs[0].add_run()
                run.add_picture(jpg_path, width=IMG_W, height=IMG_H)

                row_cap = table.add_row().cells[0]
                row_cap.text = f"{labels[key]} – Slice {slice_str}"
                if "Caption" in document.styles:
                    row_cap.paragraphs[0].style = "Caption"

                os.remove(jpg_path)

            document.add_page_break()

    out_doc = os.path.join(script_dir, "exports", f"{axis}slice_figures.docx")
    document.save(out_doc)
    print(f"Word document saved as: {os.path.abspath(out_doc)}")


# Generate
generate_report("x")
generate_report("y")
generate_report("z")

# XX
